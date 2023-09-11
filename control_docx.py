import os
import pandas as pd
from docx.shared import Inches
import streamlit as st
from docx import Document
def insert_paragraph(doc, content, style, idx):
    """
    在指定的索引位置插入一个新的段落。

    :param doc: Document 对象
    :param content: 要插入的段落的内容
    :param style: 段落的样式
    :param idx: 要插入段落的位置的索引
    :return: 新插入的段落
    """
    new_paragraph = doc.add_paragraph(content, style=style)
    p = doc.element.body[-1]
    doc.element.body.remove(p)
    doc.element.body.insert(idx, p)
    return new_paragraph

def get_section_indices(doc, title):
    section_start = None
    section_end = None
    title_found = False
    all_elements = []

    for idx, item in enumerate(doc.element.body):
        if item.tag.endswith('tbl'):
            all_elements.append({'type': 'table', 'index': idx, 'content': doc.tables[len([e for e in all_elements if e['type'] == 'table'])]})
        elif item.tag.endswith('p'):
            para = doc.paragraphs[len([e for e in all_elements if e['type'] == 'paragraph'])]
            all_elements.append({'type': 'paragraph', 'index': idx, 'content': para})

    for element in all_elements:
        if element['type'] == 'paragraph':
            current_text = element['content'].text
            current_style = element['content'].style.name
        else:  # If it's a table, we don't retrieve the text and style
            current_text = ''
            current_style = ''

        if not title_found and current_text == title:
            section_start = element['index']
            title_found = True
            continue
        elif title_found and current_style == 'Heading 1':
            section_end = element['index']
            break

    return section_start, section_end, all_elements


def delete_section_content(doc, title):
    section_start, section_end, all_elements = get_section_indices(doc, title)
    if section_start is None or section_end is None:
        return

    # Get the indices of the elements we want to delete
    to_delete_indices = [item['index'] for item in all_elements[section_start + 1:section_end]]

    # Now delete them
    for idx in reversed(to_delete_indices):  # It's important to delete from the end first
        doc.element.body.remove(doc.element.body[idx])


def extract_content_from_doc(doc: Document) -> dict:
    """从Document中提取文字和表格内容"""

    # 提取文本
    texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip() != ""]

    # 提取表格内容
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    return {
        "texts": texts,
        "tables": tables
    }





# 在doc的指定标题下添加图片
def add_images_to_section(doc, title, img_paths, width=6):
    # 查找章节标题
    paragraphs = [p for p in doc.paragraphs if p.text == title]
    if not paragraphs:
        doc.add_paragraph(title, style='Heading1')
        insertion_point = len(doc.paragraphs)
    else:
        # 使用标题文本来查找位置，而不是Paragraph对象
        insertion_point = next((i for i, p in enumerate(doc.paragraphs) if p.text == title), None) + 1
        if insertion_point is None:
            raise ValueError(f"Unable to find title '{title}' in the document.")

    # 添加图片
    for img_path in img_paths:
        new_paragraph = doc.add_paragraph()
        run = new_paragraph.add_run()
        run.add_picture(img_path, width=Inches(width))

        # 将段落移到所需的位置
        para_to_insert = new_paragraph._element
        ref_para = doc.paragraphs[insertion_point - 1]._element
        ref_para.getparent().insert(ref_para.getparent().index(ref_para) + 1, para_to_insert)
        insertion_point += 1


# 在doc的指定标题下添加文字
def add_or_update_section(doc, title, content):
    # 检查章节是否存在
    paragraphs = [p for p in doc.paragraphs if p.text == title]
    if not paragraphs:
        doc.add_paragraph(title, style='Heading1')
        insertion_point = len(doc.paragraphs) - 1
    else:
        # Instead of trying to get the index from the old Paragraph object,
        # search for the title again to get the most recent index
        insertion_point = next((i for i, p in enumerate(doc.paragraphs) if p.text == title), None) + 1
        if insertion_point is None:
            raise ValueError(f"Unable to find title '{title}' in the document.")

    # 添加新内容
    p = doc.add_paragraph(content)
    para_to_insert = p._element
    ref_para = doc.paragraphs[insertion_point - 1]._element
    ref_para.getparent().insert(ref_para.getparent().index(ref_para) + 1, para_to_insert)




# 在doc的指定标题下添加表格
def add_or_update_tables(doc, title, dataframes):
    # 检查章节是否存在
    paragraphs = [p for p in doc.paragraphs if p.text == title]
    if not paragraphs:
        doc.add_paragraph(title, style='Heading1')
        insertion_point = len(doc.paragraphs) - 1
    else:
        # Instead of trying to get the index from the old Paragraph object,
        # search for the title again to get the most recent index
        insertion_point = next((i for i, p in enumerate(doc.paragraphs) if p.text == title), None) + 1
        if insertion_point is None:
            raise ValueError(f"Unable to find title '{title}' in the document.")


    # 在指定位置后添加新表格
    for df in dataframes:
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1] + 1, style='Table Grid')

        # 添加表格内容
        table.cell(0, 0).text = df.index.name or 'Index'
        for col_num, col_name in enumerate(df.columns, start=1):
            table.cell(0, col_num).text = str(col_name)
        for row_num, row in enumerate(df.iterrows(), start=1):
            index, data = row
            table.cell(row_num, 0).text = str(index)
            for col_num, cell_value in enumerate(data, start=1):
                table.cell(row_num, col_num).text = str(cell_value)

        # 将新表格移到正确的位置
        tbl, ref_para = table._tbl, doc.paragraphs[insertion_point - 1]._element
        ref_para.getparent().insert(ref_para.getparent().index(ref_para) + 1, tbl)
        insertion_point += 1



def initialize_doc_with_titles():
    # 创建一个新的Document对象
    doc = Document()

    # 添加预设的标题
    titles = ["Title(标题) and Abstract (摘要)and Keywords (关键词)",
              "Introduction (引言)","Methods(方法)and Results(结果)",
              "Discussion (讨论)","Conclusion (结论)"]

    for title in titles:
        doc.add_paragraph(title, style='Heading1')
        # 为了增加可读性，在每个标题后加入一个空行
        doc.add_paragraph()
    return doc

