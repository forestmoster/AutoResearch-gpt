# imports
import os
import docx
import openai  # for calling the OpenAI API
from docx import Document
import split
import streamlit as st
import PyPDF2
import jieba.analyse


def extract_sentences(text_list):
    sentences = []
    for text in text_list:
        lines = text.split('\n')
        for line in lines:
            if line.strip():  # 排除空行
                sentences.append(line.strip())
    return sentences


def read_docx(file_path):
    doc = Document(file_path)
    paragraphs = []
    for paragraph in doc.paragraphs:
        paragraphs.append(paragraph.text.strip())
    return paragraphs


# 读取docx段落带字体大小和字体粗细信息
def read_doc_font(docx_file):
    doc = Document(docx_file)
    paragraphs_strings = []
    font_sizes=[]
    font_bolds=[]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:  # 跳过空字符串
            paragraphs_strings.append(text)
            for run in paragraph.runs:
                font_sizes.append(run.font.size)
                font_bolds.append(run.font.bold)
        combined_data = list(zip(paragraphs_strings, font_sizes, font_bolds))
    return combined_data


def read_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        paragraphs=extract_sentences(text)
        return paragraphs




def chunk_docx (title_long:int,MAX_TOKENS:int,BATCH_SIZE:int,uploaded_file:str):
    EMBEDDING_MODEL = "text-embedding-ada-002"
    GPT_MODEL = "gpt-3.5-turbo"
    strings = []
    text = []
    text_count = []
    titles = []
    file_name = uploaded_file.name
    # 获取文件扩展名
    try:
        paragraphs=read_doc_font(uploaded_file)
    except:
        raise ValueError("不支持的文件类型")
    for paragraph,font_size,font_bold in paragraphs:
        if len(text_count)>0and len(titles)>0and 0 < len(paragraph) <= title_long or font_bold==True :
            text_count = []
            titles = []
        if 0 < len(paragraph) <= title_long or font_bold==True:
            titles.append(paragraph)
        elif len(paragraph) > title_long:
            tags=jieba.analyse.extract_tags(paragraph,topK=5)
            text.append(paragraph)
            text_count.append(paragraph)
        if len(text) > 0:
            titles_strings= " ".join(titles)
            tags_strings = " ".join(tags)
            strings.append((file_name,titles_strings, tags_strings, text))
            text = []

    # # 对文件进行切分
    wikipedia_strings = []
    MAX_TOKENS = MAX_TOKENS
    for section in strings:
        wikipedia_strings.extend(split.split_strings_from_subsection_word(section, max_tokens=MAX_TOKENS))
    st.write(f"正在上传....{len(strings)}  sections split into {len(wikipedia_strings)} strings.")


# 导入松果数据库
    import pinecone
    # initialize connection to pinecone (get API key at app.pinecone.io)
    pinecone.init(
        api_key="b0e7c072-995c-4406-8c41-12238d626882",
        environment="us-west4-gcp"  # find next to API key in console
    )

    # check if 'openai' index already exists (only create index if not)
    if 'openai' not in pinecone.list_indexes():
        pinecone.create_index('openai', dimension=1536)
    # connect to index
    index = pinecone.Index('openai')

    from tqdm.auto import tqdm  # this is our progress bar
    current_ids = index.describe_index_stats()['total_vector_count'] - 1
    batch_size = BATCH_SIZE  # process everything in batches of 32
    for i in tqdm(range(0, len(wikipedia_strings), batch_size)):
        # set end position of batch
        i_end = min(i + batch_size, len(wikipedia_strings))
        # get batch of lines and IDs
        lines_batch = wikipedia_strings[i: i + batch_size]
        ids_batch = [str(n) for n in range(i+current_ids, i_end+current_ids)]
        # create embeddings
        res = openai.Embedding.create(input=lines_batch, engine=EMBEDDING_MODEL)
        embeds = [record['embedding'] for record in res['data']]
        # prep metadata and upsert batch
        meta = [{'text': line} for line in lines_batch]
        to_upsert = zip(ids_batch, embeds, meta)
        st.write(ids_batch, embeds, meta)
        # upsert to Pinecone
        index.upsert(vectors=to_upsert)
        last_uploaded_id = current_ids + len(ids_batch)


def chunk_pdf(title_long: int, MAX_TOKENS: int, BATCH_SIZE: int, uploaded_file: str):
    EMBEDDING_MODEL = "text-embedding-ada-002"
    GPT_MODEL = "gpt-3.5-turbo"
    file_name = uploaded_file.name
    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
    except:
        raise ValueError("不支持的文件类型")
    # file_name = os.path.basename(uploaded_file)
    num_pages = len(pdf_reader.pages)
    page_limit = 10  # 每次处理的页数限制
    start_page = 0
    count = 0

    # 导入松果数据库
    import pinecone

    # initialize connection to pinecone (get API key at app.pinecone.io)
    pinecone.init(
        api_key="b0e7c072-995c-4406-8c41-12238d626882",
        environment="us-west4-gcp"  # find next to API key in console
    )

    # check if 'openai' index already exists (only create index if not)
    if 'openai' not in pinecone.list_indexes():
        pinecone.create_index('openai', dimension=1536)

    # connect to index
    index = pinecone.Index('openai')
    last_uploaded_id=None
    while start_page < num_pages:
        count = count + 1
        print(count)
        text = []
        strings = []
        text_count = []
        titles = []
        end_page = min(start_page + page_limit, num_pages)
        for page_num in range(start_page, end_page):
            page = pdf_reader.pages[page_num]
            text.append(page.extract_text())
        # paragraphs = extract_sentences(text)
        paragraphs = text
        start_page = end_page
        # file_name = uploaded_file.name
        for paragraph in paragraphs:
            if len(text_count) > 0 and len(titles) > 0 and 0 < len(paragraph) <= title_long:
                text_count = []
                titles = []
            if 0 < len(paragraph) <= title_long:
                titles = titles.append(paragraph)
            elif len(paragraph) > title_long:
                tags = jieba.analyse.extract_tags(paragraph, topK=10)
                text.append(paragraph)
                text_count.append(paragraph)
            if len(text) > 0:
                titles_strings = " ".join(titles)
                tags_strings = " ".join(tags)
                strings.append((file_name, titles_strings,tags_strings,text))
                text = []
        wikipedia_strings = []
        MAX_TOKENS = MAX_TOKENS
        for section in strings:
            wikipedia_strings.extend(split.split_strings_from_subsection_pdf(section, max_tokens=MAX_TOKENS))
        st.write(f"正在上传....{len(strings)}  sections split into {len(wikipedia_strings)} strings.")

        from tqdm.auto import tqdm  # this is our progress bar
        if count==1:
            current_ids = index.describe_index_stats()['total_vector_count'] - 1
        else:
            current_ids=last_uploaded_id

        batch_size = BATCH_SIZE  # process everything in batches of 32
        for i in tqdm(range(0, len(wikipedia_strings), batch_size)):
            # set end position of batch
            i_end = min(i + batch_size, len(wikipedia_strings))
            # get batch of lines and IDs
            lines_batch = wikipedia_strings[i: i + batch_size]
            ids_batch = [str(n) for n in range(i+current_ids, i_end+current_ids)]
            # create embeddings
            res = openai.Embedding.create(input=lines_batch, engine=EMBEDDING_MODEL)
            embeds = [record['embedding'] for record in res['data']]
            # prep metadata and upsert batch
            meta = [{'text': line} for line in lines_batch]
            to_upsert = zip(ids_batch, embeds, meta)
            st.write(ids_batch, embeds, meta)
            # upsert to Pinecone
            index.upsert(vectors=to_upsert)
            last_uploaded_id=current_ids+len(ids_batch)


